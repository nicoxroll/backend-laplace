import os
import uuid
import time
import asyncio
from typing import List, Dict, Any, Optional, Tuple
from loguru import logger
from tqdm import tqdm
import joblib
import numpy as np
from datetime import datetime

# Solo mantener filetype:
import filetype
import json

# Procesadores de documentos
from pypdf import PdfReader
from docx import Document
import openpyxl
from bs4 import BeautifulSoup
import markdown
from unstructured.partition.auto import partition

# NLP
import nltk
import spacy
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader

# Base de datos
from sqlalchemy.orm import Session
from models import Knowledge
from database.db import SessionLocal
from db.weaviate_client import store_vectors_in_weaviate, init_schema
from db.embeddings_client import generate_embeddings

# Descargar recursos necesarios (ejecutar una vez)
def download_resources():
    """Descargar recursos necesarios para el procesamiento"""
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        spacy.cli.download("es_core_news_sm")
    except Exception as e:
        logger.warning(f"Error downloading resources: {e}")

# Función para detectar el tipo de archivo
def detect_file_type(file_path: str) -> str:
    """
    Detecta el tipo MIME de un archivo usando filetype
    """
    kind = filetype.guess(file_path)
    
    if kind is None:
        # Si filetype no detecta el tipo, usar la extensión
        import mimetypes
        mime, _ = mimetypes.guess_type(file_path)
        return mime or 'application/octet-stream'
    
    return kind.mime

# Procesadores específicos por tipo de documento
def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extrae texto de un archivo PDF"""
    pages = []
    try:
        pdf = PdfReader(file_path)
        for i, page in enumerate(tqdm(pdf.pages, desc="Procesando PDF")):
            text = page.extract_text()
            if text and text.strip():
                pages.append({
                    "content": text,
                    "metadata": {"page": i+1, "source": "pdf"}
                })
    except Exception as e:
        logger.error(f"Error procesando PDF {file_path}: {e}")
    
    return pages

def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extrae texto de un archivo Word"""
    paragraphs = []
    try:
        doc = Document(file_path)
        content = []
        for i, para in enumerate(tqdm(doc.paragraphs, desc="Procesando Word")):
            if para.text and para.text.strip():
                content.append(para.text)
            
            # Agrupar párrafos para reducir fragmentación
            if i % 5 == 4 or i == len(doc.paragraphs) - 1:
                if content:
                    paragraphs.append({
                        "content": "\n".join(content),
                        "metadata": {"section": i//5, "source": "docx"}
                    })
                    content = []
    except Exception as e:
        logger.error(f"Error procesando DOCX {file_path}: {e}")
    
    return paragraphs

def extract_text_from_excel(file_path: str) -> List[Dict[str, Any]]:
    """Extrae texto de un archivo Excel"""
    sheets = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in tqdm(wb.sheetnames, desc="Procesando Excel"):
            sheet = wb[sheet_name]
            content = []
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    content.append(row_text)
            
            if content:
                sheets.append({
                    "content": "\n".join(content),
                    "metadata": {"sheet": sheet_name, "source": "excel"}
                })
    except Exception as e:
        logger.error(f"Error procesando Excel {file_path}: {e}")
    
    return sheets

def extract_text_from_html(content: str) -> List[Dict[str, Any]]:
    """Extrae texto de contenido HTML"""
    sections = []
    try:
        soup = BeautifulSoup(content, 'html.parser')
        
        # Eliminar scripts y estilos
        for script in soup(["script", "style"]):
            script.extract()
            
        # Procesar por secciones
        for i, section in enumerate(soup.find_all(['div', 'section', 'article'])):
            text = section.get_text(strip=True, separator='\n')
            if text:
                sections.append({
                    "content": text,
                    "metadata": {"section": i, "source": "html"}
                })
        
        # Si no hay secciones, usar todo el contenido
        if not sections:
            text = soup.get_text(strip=True, separator='\n')
            if text:
                sections.append({
                    "content": text,
                    "metadata": {"section": 0, "source": "html"}
                })
    except Exception as e:
        logger.error(f"Error procesando HTML: {e}")
    
    return sections

def extract_text_with_unstructured(file_path: str) -> List[Dict[str, Any]]:
    """Extrae texto usando el paquete unstructured"""
    sections = []
    try:
        elements = partition(filename=file_path)
        
        # Agrupar elementos en secciones lógicas
        current_section = []
        current_metadata = {"source": "unstructured", "section": 0}
        
        for i, element in enumerate(elements):
            current_section.append(str(element))
            
            # Cada 10 elementos o al final, crear una sección
            if (i+1) % 10 == 0 or i == len(elements) - 1:
                if current_section:
                    sections.append({
                        "content": "\n".join(current_section),
                        "metadata": {**current_metadata, "section": len(sections)}
                    })
                    current_section = []
    
    except Exception as e:
        logger.error(f"Error procesando con unstructured {file_path}: {e}")
    
    return sections

# Dividir en chunks optimizados para embeddings
def split_into_chunks(sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Divide el texto en chunks optimizados para embeddings"""
    chunks = []
    
    # Configurar el splitter para chunks óptimos (1000-1500 caracteres)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    
    for section in tqdm(sections, desc="Dividiendo en chunks"):
        try:
            text_chunks = splitter.split_text(section["content"])
            
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():
                    chunks.append({
                        "content": chunk.strip(),
                        "metadata": {
                            **section["metadata"],
                            "chunk": i
                        }
                    })
        except Exception as e:
            logger.error(f"Error dividiendo chunks: {e}")
    
    return chunks

# Generar embeddings en paralelo
def process_chunks_with_embeddings(chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Procesa los chunks para generar embeddings en paralelo"""
    processed_chunks = []
    
    # Función para procesar un chunk individual
    def process_chunk(chunk):
        try:
            # Generar embeddings
            embedding = generate_embeddings(chunk["content"])
            
            return {
                "content": chunk["content"],
                "metadata": chunk["metadata"],
                "embedding": embedding
            }
        except Exception as e:
            logger.error(f"Error procesando chunk: {e}")
            return None
    
    # Usar joblib para paralelizar
    results = joblib.Parallel(n_jobs=-1)(
        joblib.delayed(process_chunk)(chunk) for chunk in tqdm(chunks, desc="Generando embeddings")
    )
    
    # Filtrar resultados nulos
    processed_chunks = [result for result in results if result is not None]
    
    return processed_chunks

# Función principal de procesamiento
async def process_document(file_path: str, metadata: Dict[str, Any], job_id: str):
    """
    Procesa un documento completo
    """
    start_time = time.time()
    
    try:
        # Verificar que el archivo exista
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        # 1. Detectar tipo de archivo
        mime_type = detect_file_type(file_path)
        
        # Actualizar estado
        update_processing_status(job_id, {
            "status": "processing", 
            "progress": 0.1, 
            "message": f"Detectado archivo: {mime_type}"
        })
        
        # 2. Extraer texto según el tipo
        sections = []
        
        if "pdf" in mime_type:
            sections = extract_text_from_pdf(file_path)
        elif "word" in mime_type or "docx" in mime_type:
            sections = extract_text_from_docx(file_path)
        elif "excel" in mime_type or "xlsx" in mime_type:
            sections = extract_text_from_excel(file_path)
        elif "html" in mime_type:
            with open(file_path, 'r', encoding='utf-8') as f:
                sections = extract_text_from_html(f.read())
        else:
            # Usar unstructured como fallback
            sections = extract_text_with_unstructured(file_path)
        
        if not sections:
            raise ValueError(f"No se pudo extraer texto del archivo {file_path}")
        
        # Actualizar estado
        update_processing_status(job_id, {
            "status": "processing", 
            "progress": 0.3, 
            "message": f"Texto extraído: {len(sections)} secciones"
        })
        
        # 3. Dividir en chunks optimizados
        chunks = split_into_chunks(sections)
        
        update_processing_status(job_id, {
            "status": "processing", 
            "progress": 0.5, 
            "message": f"Generando embeddings para {len(chunks)} chunks"
        })
        
        # 4. Procesar chunks y generar embeddings
        processed_chunks = process_chunks_with_embeddings(chunks)
        
        # 5. Almacenar en Weaviate
        update_processing_status(job_id, {
            "status": "processing", 
            "progress": 0.8, 
            "message": "Guardando vectores en Weaviate"
        })
        
        vector_ids = store_vectors_in_weaviate(processed_chunks, {
            "user_id": metadata["user_id"],
            "filename": metadata["filename"],
            "job_id": job_id,
            "content_type": mime_type,
            "processed_at": datetime.now().isoformat()
        })
        
        # 6. Crear registro en la base de datos SQL
        update_processing_status(job_id, {
            "status": "processing", 
            "progress": 0.9, 
            "message": "Registrando conocimiento en la base de datos"
        })
        
        db = SessionLocal()
        try:
            # Crear nuevo Knowledge con vector_ids
            knowledge = Knowledge(
                user_id=metadata["user_id"],
                name=metadata["filename"],
                description=f"Archivo procesado: {metadata['filename']}",
                content_hash=job_id,
                vector_ids=vector_ids,
                base_id=metadata.get("base_id")
            )
            db.add(knowledge)
            db.commit()
            
            # Actualizar estado completado con knowledge_id
            knowledge_id = knowledge.id
            update_processing_status(job_id, {
                "status": "completed",
                "progress": 1.0,
                "message": "Procesamiento completado con éxito",
                "completed_at": datetime.now().isoformat(),
                "knowledge_id": knowledge_id
            })
            
            # Retornar resultado
            elapsed_time = time.time() - start_time
            logger.info(f"Documento procesado en {elapsed_time:.2f} segundos, ID: {knowledge_id}")
            return {
                "status": "completed",
                "knowledge_id": knowledge_id,
                "vector_ids": vector_ids,
                "chunks": len(processed_chunks),
                "elapsed_time": elapsed_time
            }
            
        except Exception as e:
            db.rollback()
            raise e
        finally:
            db.close()
        
    except Exception as e:
        logger.error(f"Error en procesamiento: {str(e)}")
        update_processing_status(job_id, {
            "status": "failed",
            "message": f"Error: {str(e)}"
        })
        raise e

# Gestión de estados con Redis
import os
import redis
from dotenv import load_dotenv

# Cargar variables de entorno
load_dotenv()

# Usar la URL de Redis desde variables de entorno
REDIS_URL = os.getenv("REDIS_URL", "redis://redis:6379/0")

try:
    # Usar REDIS_URL en lugar de localhost:6379
    redis_client = redis.from_url(REDIS_URL)
    redis_client.ping()  # Verificar la conexión
    logger.info(f"Conectado a Redis correctamente en {REDIS_URL}")
except Exception as e:
    logger.warning(f"No se pudo conectar a Redis: {e}")
    # Fallback a diccionario en memoria
    in_memory_status = {}
    
def update_processing_status(job_id: str, status_data: Dict[str, Any]):
    """Actualiza el estado de procesamiento en Redis"""
    try:
        # Asegurarse de que status_data tiene todos los campos necesarios
        if "updated_at" not in status_data:
            status_data["updated_at"] = datetime.now().isoformat()
            
        if "job_id" not in status_data:
            status_data["job_id"] = job_id
            
        # Intentar guardar en Redis
        try:
            redis_client.setex(
                f"job_status:{job_id}", 
                86400,  # TTL: 24 horas
                json.dumps(status_data)
            )
        except NameError:
            # Fallback a memoria
            in_memory_status[job_id] = status_data
            
        logger.debug(f"Estado actualizado para job {job_id}: {status_data.get('status')} {status_data.get('progress')}")
        
    except Exception as e:
        logger.error(f"Error actualizando estado para job {job_id}: {e}")

def get_processing_status(job_id: str) -> Dict[str, Any]:
    """Obtiene el estado actual de un trabajo desde Redis"""
    try:
        # Intentar obtener desde Redis
        try:
            status_json = redis_client.get(f"job_status:{job_id}")
            if status_json:
                return json.loads(status_json)
        except NameError:
            # Fallback a memoria
            if job_id in in_memory_status:
                return in_memory_status[job_id]
                
        return None
    except Exception as e:
        logger.error(f"Error obteniendo estado para job {job_id}: {e}")
        return None

# Añadir función específica para procesar repositorios
async def process_repository_json(file_path: str, job_id: str, user_id: str, metadata: dict):
    """
    Procesa un archivo JSON de repositorio y lo vectoriza para Weaviate
    """
    logger.info(f"Procesando repositorio desde JSON: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            repo_data = json.load(f)
        
        # Extraer contenido del repositorio
        vectors = []
        
        # Recorrer la estructura del repositorio
        for file_item in repo_data.get('files', []):
            file_path = file_item.get('path', '')
            file_content = file_item.get('content', '')
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Saltar archivos binarios o sin contenido
            if not file_content or is_binary_content(file_extension):
                continue
                
            # Chunking del contenido del archivo
            chunks = create_chunks(file_content, 1000, 200)
            
            # Generar embeddings para los chunks
            embeddings = await generate_embeddings(chunks)
            
            # Crear vectores para cada chunk
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vectors.append({
                    "content": chunk,
                    "embedding": embedding,
                    "metadata": {
                        "file_path": file_path,
                        "chunk_index": i
                    },
                    "batch_id": len(vectors)  # Índice único para cada vector
                })
        
        # Metadatos para Weaviate
        weaviate_metadata = {
            "user_id": user_id,
            "filename": metadata.get("filename", "repository.json"),
            "job_id": job_id,
            "content_type": "repository",
            "processed_at": datetime.now().isoformat()
        }
        
        # Almacenar vectores en Weaviate
        vector_ids = store_vectors_in_weaviate(vectors, weaviate_metadata)
        logger.info(f"Repositorio indexado exitosamente: {len(vector_ids)} chunks")
        
        return {
            "status": "completed",
            "vectors": len(vector_ids),
            "vector_ids": vector_ids
        }
        
    except Exception as e:
        logger.error(f"Error procesando repositorio JSON: {str(e)}")
        return {
            "status": "failed",
            "error": str(e)
        }